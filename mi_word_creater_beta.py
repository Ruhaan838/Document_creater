from docx import Document
from tkinter import *
from tkinter import filedialog,messagebox
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor,Inches,Pt
import customtkinter as ctk
#define the globel variabel is help to store the location
file_loacation = ""
error = []
#creating the docx file
class work:
    def __init__(self):
        self.problem_open1 =[]
        self.output_open1  =[]
        self.code_open1    =[]
        self.Input_open1   =[]
        
    #make function of the image open and text open and loction_folder
    def txt_open(self):
        problem_open = filedialog.askopenfilename(
            title="Select a file",
            initialdir="/home/user",
            filetypes=[("Text files", "*.txt")]
        )
        self.problem_open1.append(problem_open)
        
    def image_open(self):
        output_open = filedialog.askopenfilename(
            title="Select a file",
            initialdir="/home/user",
            filetypes=[("PNG files", "*.png")]
        )
        self.output_open1.append(output_open)
    def image_open_1(self):
        input_png_input = filedialog.askopenfilename(
            title="Select a file",
            initialdir="/home/user",
            filetypes=[("PNG files", "*.png")]
        )
        self.Input_open1.append(input_png_input)    
    def txt_open_1(self):
        code_open  = filedialog.askopenfilename(
            title="Select a file",
            initialdir="/home/user",
            filetypes=[("Text files", "*.txt")]
        )
        self.code_open1.append(code_open)
    def docx_file(self):
        doc = Document()
        
        #header work
        for section in doc.sections:
            header = section.header
            footer = section.footer    
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_run = header_paragraph.add_run()
        rollno = roll_no_input.get()       
        header_run.add_text(f"Enrollment No:{rollno}")
        header_run.add_text("\t\t Microprocessor and Interfacing (BTCO13304)")
        
        #header formating
        header_run.bold = True
        header_run.font.size = Pt(11)
        
        #footer work
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_run = footer_paragraph.add_run("2023-24/BE_II_CO_DIV-I/Sem-3/MI        Computer Engg. Deptt., SCET, Surat \t Page NO:   ")
        
        #some formating on the fotter
        footer_run.font.size = Pt(11)
        
        #page no:
        page_num_field = OxmlElement('w:fldSimple')
        page_num_field.set(qn('w:instr'), 'PAGE')
        page_num_run = footer_paragraph.add_run()
        page_num_run._r.append(page_num_field)
        page_num_run.font.name = 'Arial'
        page_num_run.font.size = Pt(11)
        
        #the main file work
        no_of_p = practical_input.get()
        page_num = int(no_of_p)
        for page_no in range(page_num):
            
            #titel for practical 
            practical_name = doc.add_heading(f"Practical {page_no+1}",level=0)
            
            #formating of the Practical
            practical_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for colr in practical_name.runs:
                colr.font.color.rgb = RGBColor(255,183,3)
                
#problem statement    
            subtitle = doc.add_paragraph('Problem Statement:')
            
            #formating of the problem Statement
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                
            #read the text file for the problem statement   
            with open(self.problem_open1[page_no],'r') as Problem_file:
                para1 = Problem_file.read()
            para = doc.add_paragraph(para1)
            
            #formating of the text file-1
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                
#code Text file
            subtitle = doc.add_paragraph('Code:')
            
            #formating of subtitle
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'
                
            #read the code text file
            with open(self.code_open1[page_no],'r') as code_file:
                para1 = code_file.read()
            para = doc.add_paragraph(para1)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'
            doc.add_page_break()
#input image file               
            subtitle = doc.add_paragraph('Input:')
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'
            doc.add_picture(self.Input_open1[page_no],width=Inches(7),height=Inches(3))     
#output image file                
            subtitle = doc.add_paragraph('Output:')
            for run in subtitle.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'   
            doc.add_picture(self.output_open1[page_no],width=Inches(7),height=Inches(3))
            if(page_no == page_num-1):
                continue
            doc.add_page_break()
        save_docx = file_loacation + "/output.docx"
        error.append(save_docx)                                              
        doc.save(save_docx)
    
def location_folder():
    global file_loacation
    file_loacation_input = filedialog.askdirectory(
        title="Select a folder",
        initialdir="/home/user"
    )
    file_loacation = file_loacation_input 
def sclect_file():
    global current_page
    no_of_p = practical_input.get()
    page_num = int(no_of_p)
    if (current_page < page_num):
        current_page += 1
        massage = Label(dielog_box,text=f"Select the file for the Practical :-{current_page} ",foreground="White",background='#242424').place(x=40,y=100)
        problem_text.config(text="Select the Problem Stetment file",foreground="White",background='#242424')
        code_text.config(text="Select the Code Text file:",foreground="White",background='#242424')
        input_png.config(text="Select the Input png file:",foreground="White",background='#242424')
        output_png.config(text="Select the Output png file:",foreground="White",background='#242424')
        show_buttons()
        hide_buttons1()
    elif (page_num == 0 or page_num == None):
        messagebox.showerror("Opps !!", "Enter The Practical Number")
    elif (current_page == page_num):
        complete.config(text="Selection is completed!",foreground="White",background='#242424')
        document_save.config(text="Select folder for store the .docx file:",foreground="White",background='#242424')     
        hide_buttons()
        show_buttons1()
def submit():
    # Check if all required files are selected
    if not work_instance.problem_open1 or not work_instance.code_open1 or not work_instance.Input_open1 or not work_instance.output_open1:
        messagebox.showerror("Error", "Please select all required files.")
        return

    # Check if a folder for saving the .docx file is selected
    if not file_loacation:
        messagebox.showerror("Error", "Please select a folder to save the .docx file.")
        return

    # Proceed to create the .docx file
    try:
        work_instance.docx_file()
        messagebox.showinfo("Success", f"Your file is saved in {error} location")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")
def About():
    ab =Tk()
    ab.title("About App !")
    ab.geometry("500x200")
    ab.resizable(False, False)
    intro = Label(master=ab, text="""
This app is basically made for the help for creating the Document File 
only You need to 
(1) Enter Your Enrollment NO
(2)Add The problem statement in Form of the text file
(3) Add The Code Of the Given Input in Form of the text file
(4)Add the Photo or Sceenshot of your output 
(5)Select the location of the Folder Which You store your docx file
(6) if You can't select any one of them it will give you error !!""",)
    intro.pack()               
#butten showing and hiding
def hide_buttons():
    problem_text.place_forget()
    problem_text_input.place_forget()
    code_text.place_forget()
    code_text_input.place_forget()
    input_png.place_forget()
    input_png_input.place_forget()
    output_png.place_forget()
    output_png_input.place_forget()
def show_buttons():
    problem_text.place(x=40, y=150)
    problem_text_input.place(x=300, y=150)
    code_text.place(x=40, y=190)
    code_text_input.place(x=300, y=190)
    input_png.place(x=40,y=230)
    input_png_input.place(x=300,y=230)
    output_png.place(x=40, y=270)
    output_png_input.place(x=300, y=270)
def show_buttons1():
    document_save.place(x=40,y=220)
    document_save_input.place(x=300,y=220) 
def hide_buttons1():
    document_save.place_forget()
    document_save.place_forget()
def done():
    messagebox.showinfo("successfully", f"Your file is save in {error} location")                                     
dielog_box = ctk.CTk()
   
#define the size of the dilog box
dielog_box.geometry("500x380")
dielog_box.resizable(False, False)
dielog_box.configure(bg='#353535')
#give the titel
dielog_box.title("Docx Creater")
#about button
About_Button = ctk.CTkButton(master=dielog_box, text="Help !", width=3, command=About)
About_Button.place(x=2, y=0)
#the massage
custom_font = ("Borneox", 18,'bold')
intro = Label(dielog_box, text="Fill This Instruction",foreground="White",background='#242424',font=custom_font)

#for enrollment no 
roll_no = Label(dielog_box, text="Enrollment No:",foreground="White",background='#242424')
roll_no.place(x=40,y=40)
roll_no_input = ctk.CTkEntry(dielog_box, width=150)
roll_no_input.place(x=300, y=40)

work_instance = work()
current_page = 0

#for asking practical no
practical = Label(dielog_box, text="Which amount of practical you have:",foreground="White",background='#242424')
practical.place(x=40,y=70)
practical_input = ctk.CTkComboBox(dielog_box, values=['0','1','2','3','4','5'])
practical_input.place(x=300,y=70)

#complete the selecttion
complete = Label(master=dielog_box,text="",foreground="White",background='#242424')
complete.place(x=180,y=250)

#for secelt the problem stetment file as .txt
problem_text = Label(dielog_box,text="")
problem_text_input =ctk.CTkButton(master=dielog_box,text="Choose the (.txt) file",command=work_instance.txt_open)

#for select the code file as .txt
code_text = Label(dielog_box,text="")   
code_text_input = ctk.CTkButton(master=dielog_box, text="Choose the (.txt) file", command=work_instance.txt_open_1)

#for select the photo for the input png
input_png = Label(dielog_box,text="")
input_png_input = ctk.CTkButton(master=dielog_box,text="Choose the (.png) file",command=work_instance.image_open_1)
#for select the photo for the output png
output_png = Label(dielog_box,text="")
output_png_input = ctk.CTkButton(master=dielog_box,text="Choose the (.png) file",command=work_instance.image_open)    

#for stoting the .docx file
document_save = Label(dielog_box,text="")
document_save_input = ctk.CTkButton(master=dielog_box, text="Select the folder", command=location_folder)

#next button
next_button = ctk.CTkButton(master=dielog_box,text ="Next",command=sclect_file)
next_button.place(x=190,y=330)

#submit button
submit_button = ctk.CTkButton(master=dielog_box, border_width=2,text = "Submit",command =submit)
submit_button.place(x=340,y=330)            

#exit button
exit_button = ctk.CTkButton(master=dielog_box,text="Exit",command=dielog_box.destroy)
exit_button.place(x=40,y=330)


intro.pack()
dielog_box.mainloop()
